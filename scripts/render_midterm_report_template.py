from __future__ import annotations

import argparse
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt


WORDCONV = Path(r"C:\Program Files\Microsoft Office\root\Office16\Wordconv.exe")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template-doc", required=True)
    parser.add_argument("--template-docx", required=True)
    parser.add_argument("--markdown", required=True)
    parser.add_argument("--output-docx", required=True)
    parser.add_argument("--student-name", required=True)
    parser.add_argument("--student-id", required=True)
    parser.add_argument("--class-name", required=True)
    parser.add_argument("--phone", required=True)
    parser.add_argument("--email", required=True)
    parser.add_argument("--internship", default="")
    parser.add_argument("--thesis-title", required=True)
    parser.add_argument("--report-date", required=True)
    return parser.parse_args()


def convert_doc_to_docx(template_doc: Path, template_docx: Path) -> None:
    template_docx.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [str(WORDCONV), "-oice", "-nme", str(template_doc), str(template_docx)],
        check=True,
    )
    if not template_docx.exists():
        raise FileNotFoundError(f"Converted template not found: {template_docx}")


def get_sections(markdown_path: Path) -> list[tuple[str, str]]:
    raw = markdown_path.read_text(encoding="utf-8")
    matches = re.findall(r"(?ms)^##\s+(.+?)\n(.*?)(?=^##\s+|\Z)", raw)
    if len(matches) < 5:
        raise ValueError(f"Expected at least 5 sections, got {len(matches)}")
    return [(title.strip(), body.strip()) for title, body in matches]


def markdown_body_to_paragraphs(text: str) -> list[str]:
    normalized = text.replace("\r\n", "\n")
    paragraphs: list[str] = []
    buffer: list[str] = []

    for raw_line in normalized.split("\n"):
        line = raw_line.strip()
        if not line:
            if buffer:
                paragraphs.append(" ".join(buffer).strip())
                buffer = []
            continue

        numbered = re.sub(r"^\d+\.\s+", "", line)
        bulleted = re.sub(r"^-+\s+", "", numbered)
        if line != numbered or line != bulleted:
            if buffer:
                paragraphs.append(" ".join(buffer).strip())
                buffer = []
            paragraphs.append(bulleted.strip())
            continue

        buffer.append(line)

    if buffer:
        paragraphs.append(" ".join(buffer).strip())

    return paragraphs


def set_run_font(run, size_pt: float) -> None:
    run.font.name = "宋体"
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def fill_single_paragraph(paragraph, text: str, size_pt: float = 12, center: bool = False) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size_pt)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def fill_cell_with_paragraphs(cell, paragraphs: list[str], size_pt: float = 12) -> None:
    cell.text = ""
    if not paragraphs:
        paragraphs = [""]

    first = cell.paragraphs[0]
    first.text = ""
    for idx, text in enumerate(paragraphs):
        paragraph = first if idx == 0 else cell.add_paragraph()
        paragraph.text = ""
        run = paragraph.add_run(text)
        set_run_font(run, size_pt)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)


def main() -> None:
    args = parse_args()

    template_doc = Path(args.template_doc).resolve()
    template_docx = Path(args.template_docx).resolve()
    output_docx = Path(args.output_docx).resolve()
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    convert_doc_to_docx(template_doc, template_docx)

    sections = get_sections(Path(args.markdown).resolve())
    summary = markdown_body_to_paragraphs(sections[1][1])
    plan = markdown_body_to_paragraphs(sections[2][1])
    problems = markdown_body_to_paragraphs(sections[3][1])
    other = markdown_body_to_paragraphs(sections[4][1])

    doc = Document(str(template_docx))

    fill_single_paragraph(doc.paragraphs[1], args.report_date, size_pt=10.5, center=True)

    table = doc.tables[0]
    fill_single_paragraph(table.rows[0].cells[1].paragraphs[0], args.student_id)
    fill_single_paragraph(table.rows[0].cells[3].paragraphs[0], args.student_name)
    fill_single_paragraph(table.rows[0].cells[5].paragraphs[0], args.class_name)
    fill_single_paragraph(table.rows[1].cells[1].paragraphs[0], args.phone)
    fill_single_paragraph(table.rows[1].cells[4].paragraphs[0], args.email)
    fill_single_paragraph(table.rows[2].cells[1].paragraphs[0], args.internship)
    fill_single_paragraph(table.rows[2].cells[4].paragraphs[0], args.thesis_title)

    fill_cell_with_paragraphs(table.rows[3].cells[1], summary)
    fill_cell_with_paragraphs(table.rows[4].cells[1], plan)
    fill_cell_with_paragraphs(table.rows[5].cells[1], problems)
    fill_cell_with_paragraphs(table.rows[6].cells[1], other)

    doc.save(str(output_docx))
    print(output_docx)


if __name__ == "__main__":
    main()
